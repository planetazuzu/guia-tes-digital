#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documento Word del Manual TES Digital
FASE 3: Generación de documento Word con hipervínculos internos
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Length

# Configuración
BASE_DIR = Path("/home/planetazuzu/protocolo-r-pido/manual-tes/TES_Manual_Digital")
OUTPUT_FILE = Path("/home/planetazuzu/protocolo-r-pido/MANUAL_TES_DIGITAL_COMPLETO.docx")

# Mapeo de archivos según estructura jerárquica (rutas relativas desde BASE_DIR)
ESTRUCTURA_ARCHIVOS = {
    "PARTE I: FUNDAMENTOS Y EVALUACIÓN INICIAL": {
        "BLOQUE 0 - Fundamentos de Emergencias Prehospitalarias": [
            "./BLOQUE_0_FUNDAMENTOS/BLOQUE_00_0_FUNDAMENTOS_EMERGENCIAS.md"
        ],
        "BLOQUE 1 - Procedimientos Básicos": [
            "./BLOQUE_1_PROCEDIMIENTOS_BASICOS/BLOQUE_01_1_CONSTANTES_VITALES.md",
            "./BLOQUE_1_PROCEDIMIENTOS_BASICOS/BLOQUE_01_2_ABCDE_OPERATIVO.md",
            "./BLOQUE_1_PROCEDIMIENTOS_BASICOS/BLOQUE_01_3_GLASGOW_OPERATIVO.md",
            "./BLOQUE_1_PROCEDIMIENTOS_BASICOS/BLOQUE_01_4_TRIAGE_START.md"
        ]
    },
    "PARTE II: SOPORTE VITAL Y PROCEDIMIENTOS CRÍTICOS": {
        "BLOQUE 4 - Soporte Vital Básico y RCP": [
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_0_ACCESO_VASCULAR_BASICO.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_0B_RECONOCIMIENTO_PCR.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_1_RCP_ADULTOS.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_2_RCP_PEDIATRIA.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_3_RCP_LACTANTES.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_4_USO_DESA.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_5_RCP_DOS_INTERVINIENTES.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_6_OVACE_ADULTOS.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_7_OVACE_PEDIATRIA.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_8_OVACE_LACTANTES.md",
            "./BLOQUE_4_SOPORTE_VITAL_BASICO_Y_RCP/BLOQUE_04_9_POSICION_LATERAL_SEGURIDAD.md"
        ],
        "BLOQUE 9 - Medicina de Emergencias Aplicada": [
            "./BLOQUE_9_MEDICINA_EMERGENCIAS_APLICADA/BLOQUE_09_0_MEDICINA_EMERGENCIAS_APLICADA.md"
        ]
    },
    "PARTE III: MATERIAL Y EQUIPAMIENTO": {
        "BLOQUE 2 - Material e Inmovilización": [
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_0_ANATOMIA_OPERATIVA.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_2_INMOVILIZACION_MANUAL.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_3_COLLARIN_CERVICAL.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_4_CAMILLA_CUCHARA.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_5_TABLERO_ESPINAL.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_6_COLCHON_VACIO.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_7_EXTRICACION_MOVIMIENTOS_BLOQUE.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_8_TRANSFERENCIAS_MOVILIZACION.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_9_ERRORES_CRITICOS.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_10_FERULAS.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_11_CINTURON_PELVICO.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_12_FERULA_TRACCION.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_13_CAMILLAS_SILLAS_EVACUACION.md",
            "./BLOQUE_2_MATERIAL_E_INMOVILIZACION/BLOQUE_02_X_INVENTARIO_MATERIAL.md"
        ],
        "BLOQUE 3 - Material Sanitario y Oxigenoterapia": [
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_0_OXIGENOTERAPIA_BASICA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_0_OXIGENOTERAPIA_FUNDAMENTOS.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_1_DISPOSITIVOS_OXIGENOTERAPIA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_1_VENTILACION_BOLSA_MASCARILLA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_2_ASPIRACION.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_2_CANULA_OROFARINGEA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_3_BVM.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_4_CANULAS.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_5_ORGANIZACION_MALETIN.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_6_CONTROL_HEMORRAGIAS.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_7_QUEMADURAS.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_8_HERIDAS_VENDAJES.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_9_EXPOSICION_AISLAMIENTO_TERMICO.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_10_MONITORIZACION_BASICA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_11_GLUCOMETRO.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_12_TERMOMETRIA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_13_CONFORT_DOLOR.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_14_BIOSEGURIDAD_DESCONTAMINACION.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_15_GESTION_MATERIAL_ESCENA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_16_COMUNICACION_OPERATIVA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_17_SENALIZACION_ILUMINACION.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_18_DOCUMENTACION_OPERATIVA.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_99_CIERRE_BLOQUE_3.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_X_INVENTARIO_MATERIAL_SANITARIO.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_X2_MALETIN_CURAS.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_X3_BOLSA_MONITORIZACION.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_X4_INVENTARIO_GLOBAL.md",
            "./BLOQUE_3_MATERIAL_SANITARIO_Y_OXIGENOTERAPIA/BLOQUE_03_X5_CHECKLIST_MAESTRO.md"
        ]
    },
    "PARTE IV: FARMACOLOGÍA Y MEDICAMENTOS": {
        "BLOQUE 6 - Farmacología y Vademécum Operativo": [
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_0_PRINCIPIOS_ADMINISTRACION_FARMACOS.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_1_VADEMECUM_OPERATIVO.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_2_OXIGENO_ADMINISTRACION_Y_SEGURIDAD.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_3_ADRENALINA_USO_ANAFILAXIA_Y_RCP.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_4_ASPIRINA_USO_SCA.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_5_GLUCAGON_USO_HIPOGLUCEMIA.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_6_SALBUTAMOL_USO_CRISIS_ASMATICA.md",
            "./BLOQUE_6_FARMACOLOGIA/BLOQUE_06_7_ABREVIATURAS_TERMINOLOGIA_FARMACOLOGICA.md"
        ]
    },
    "PARTE V: PROTOCOLOS Y GESTIÓN OPERATIVA": {
        "BLOQUE 5 - Protocolos Transtelefónicos": [
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_0_INTRODUCCION_PROTOCOLOS_TRANSTELEFONICOS.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_0_PROTOCOLOS_EMERGENCIAS_ESPECIFICAS.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_1_PCR_TRANSTELEFONICA.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_2_OVACE_TRANSTELEFONICA.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_3_SCA_TRANSTELEFONICO.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_4_ICTUS_TRANSTELEFONICO.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_5_ANAFILAXIA_TRANSTELEFONICA.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_6_CRISIS_ASMATICA_TRANSTELEFONICA.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_7_HIPOGLUCEMIA_TRANSTELEFONICA.md",
            "./BLOQUE_5_PROTOCOLOS_TRANSTELEFONICOS/BLOQUE_05_8_COMUNICACION_COORDINADOR.md"
        ],
        "BLOQUE 8 - Gestión Operativa y Documentación": [
            "./BLOQUE_8_GESTION_OPERATIVA_Y_DOCUMENTACION/BLOQUE_08_0_INTRODUCCION_GESTION_OPERATIVA.md",
            "./BLOQUE_8_GESTION_OPERATIVA_Y_DOCUMENTACION/BLOQUE_08_1_DOCUMENTACION_CLINICA_PREHOSPITALARIA.md",
            "./BLOQUE_8_GESTION_OPERATIVA_Y_DOCUMENTACION/BLOQUE_08_2_COORDINACION_Y_COMUNICACION_OPERATIVA.md",
            "./BLOQUE_8_GESTION_OPERATIVA_Y_DOCUMENTACION/BLOQUE_08_3_GESTION_RECURSOS_Y_MATERIAL.md",
            "./BLOQUE_8_GESTION_OPERATIVA_Y_DOCUMENTACION/BLOQUE_08_4_CALIDAD_Y_MEJORA_CONTINUA.md"
        ]
    },
    "PARTE VI: CONDUCCIÓN Y SEGURIDAD VIAL": {
        "BLOQUE 7 - Conducción y Seguridad Vial": [
            "./BLOQUE_7_CONDUCCION_Y_SEGURIDAD_VIAL/BLOQUE_07_0_FUNDAMENTOS_CONDUCCION_URGENCIAS.md",
            "./BLOQUE_7_CONDUCCION_Y_SEGURIDAD_VIAL/BLOQUE_07_1_USO_LUCES_Y_SIRENA.md",
            "./BLOQUE_7_CONDUCCION_Y_SEGURIDAD_VIAL/BLOQUE_07_2_TECNICAS_CONDUCCION_EMERGENCIAS.md",
            "./BLOQUE_7_CONDUCCION_Y_SEGURIDAD_VIAL/BLOQUE_07_3_SEGURIDAD_VIAL_Y_PREVENCION_ACCIDENTES.md",
            "./BLOQUE_7_CONDUCCION_Y_SEGURIDAD_VIAL/BLOQUE_07_4_GESTION_RUTAS_Y_NAVEGACION.md",
            "./BLOQUE_7_CONDUCCION_Y_SEGURIDAD_VIAL/BLOQUE_07_5_PROTOCOLOS_SEGURIDAD_EN_ESCENA.md"
        ]
    },
    "PARTE VII: SITUACIONES ESPECIALES Y TRAUMA": {
        "BLOQUE 10 - Situaciones Especiales y Protocolos Avanzados": [
            "./BLOQUE_10_SITUACIONES_ESPECIALES/BLOQUE_10_0_SITUACIONES_ESPECIALES.md"
        ],
        "BLOQUE 11 - Protocolos de Trauma y Escenarios de Riesgo": [
            "./BLOQUE_11_PROTOCOLOS_TRAUMA/BLOQUE_11_0_PROTOCOLOS_TRAUMA.md"
        ]
    },
    "PARTE VIII: HABILIDADES PROFESIONALES": {
        "BLOQUE 12 - Marco Legal, Ético y Profesional del TES": [
            "./BLOQUE_12_MARCO_LEGAL_ETICO_PROFESIONAL/BLOQUE_12_0_MARCO_LEGAL_ETICO_PROFESIONAL.md"
        ],
        "BLOQUE 13 - Comunicación y Relación con el Paciente": [
            "./BLOQUE_13_COMUNICACION_RELACION_PACIENTE/BLOQUE_13_0_COMUNICACION_RELACION_PACIENTE.md"
        ],
        "BLOQUE 14 - Seguridad Personal y Salud del TES": [
            "./BLOQUE_14_SEGURIDAD_PERSONAL_SALUD_TES/BLOQUE_14_0_SEGURIDAD_PERSONAL_SALUD_TES.md"
        ]
    }
}

def sanitize_bookmark(text):
    """Sanitiza texto para usar como bookmark"""
    # Elimina caracteres especiales y espacios
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[-\s]+', '_', text)
    return text[:50]  # Limita longitud

def process_markdown_content(doc, content, file_name):
    """Procesa contenido Markdown y lo añade al documento Word con formato compacto"""
    lines = content.split('\n')
    in_code_block = False
    last_was_empty = False
    
    for line in lines:
        # Detectar bloques de código
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line, style='Normal')
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            last_was_empty = False
            continue
        
        # Saltos de página explícitos - eliminar completamente
        if line.strip() == '---' and len(line.strip()) == 3:
            continue
        
        # Títulos
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            # Limpiar formato markdown del título
            title_text = re.sub(r'\*\*(.+?)\*\*', r'\1', title_text)
            title_text = re.sub(r'\*(.+?)\*', r'\1', title_text)
            
            if level == 1:
                p = doc.add_heading(title_text, level=1)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
            elif level == 2:
                p = doc.add_heading(title_text, level=2)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
            elif level == 3:
                p = doc.add_heading(title_text, level=3)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
            elif level == 4:
                p = doc.add_heading(title_text, level=4)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(2)
            elif level == 5:
                p = doc.add_heading(title_text, level=5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(1)
            else:
                p = doc.add_heading(title_text, level=6)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(1)
            
            last_was_empty = False
            continue
        
        # Listas con guiones
        if line.strip().startswith('-') or line.strip().startswith('*'):
            list_text = line.strip().lstrip('-*').strip()
            
            # Procesar formato dentro de la lista
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            process_text_formatting(p, list_text)
            last_was_empty = False
            continue
        
        # Listas numeradas
        match = re.match(r'^(\d+)\.\s*(.+)', line.strip())
        if match:
            list_text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            process_text_formatting(p, list_text)
            last_was_empty = False
            continue
        
        # Líneas vacías - máximo una seguida
        if not line.strip():
            if not last_was_empty:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                last_was_empty = True
            continue
        
        # Texto normal
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        process_text_formatting(p, line.strip())
        last_was_empty = False

def process_text_formatting(paragraph, text):
    """Procesa formato Markdown en texto (negritas, cursivas, etc.)"""
    # Patrón para negritas **texto**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Negrita
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Cursiva (solo si no es negrita)
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.strip():
            # Texto normal
            paragraph.add_run(part)

def process_markdown_file(file_path):
    """Lee un archivo Markdown"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"    ⚠️  Error leyendo {file_path}: {e}")
        return None

def main():
    """Función principal"""
    print("=" * 70)
    print("GENERACIÓN DE DOCUMENTO WORD - MANUAL TES DIGITAL")
    print("FASE 3: Documento Word con hipervínculos internos")
    print("=" * 70)
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    
    # Configurar estilos de títulos
    for i in range(1, 7):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(max(16 - i, 12))  # Título 1: 16pt, Título 2: 15pt, etc.
        heading_style.paragraph_format.line_spacing = 1.15
        if i == 1:
            heading_style.paragraph_format.space_before = Pt(6)
            heading_style.paragraph_format.space_after = Pt(3)
        elif i == 2:
            heading_style.paragraph_format.space_before = Pt(6)
            heading_style.paragraph_format.space_after = Pt(3)
        else:
            heading_style.paragraph_format.space_before = Pt(4)
            heading_style.paragraph_format.space_after = Pt(2)
    
    # Configurar estilo de listas
    for list_style_name in ['List Bullet', 'List Number']:
        if list_style_name in doc.styles:
            list_style = doc.styles[list_style_name]
            list_style.font.size = Pt(11)
            list_style.paragraph_format.line_spacing = 1.15
            list_style.paragraph_format.space_before = Pt(0)
            list_style.paragraph_format.space_after = Pt(0)
    
    # Configurar márgenes (compactos: 2cm = ~0.79 pulgadas)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
    
    # PORTADA (compacta)
    print("\n📄 Creando portada...")
    title = doc.add_heading('MANUAL TES DIGITAL', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph('Versión 1.0')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    
    p = doc.add_paragraph('Fecha: 2024-12-15')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    
    p = doc.add_paragraph('Tipo: Manual Operativo Completo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ÍNDICE GENERAL (compacto)
    print("📑 Creando índice general...")
    doc.add_heading('ÍNDICE GENERAL', level=1)
    
    for parte, bloques in ESTRUCTURA_ARCHIVOS.items():
        p = doc.add_paragraph(parte, style='Normal')
        p.style.font.bold = True
        p.style.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        
        for bloque, archivos in bloques.items():
            p = doc.add_paragraph(f"  {bloque}", style='Normal')
            p.style.font.size = Pt(11)
            p.style.font.italic = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
    
    doc.add_page_break()
    
    # PROCESAR CADA PARTE
    total_archivos = sum(len(archivos) for bloques in ESTRUCTURA_ARCHIVOS.values() for archivos in bloques.values())
    archivo_actual = 0
    
    partes_list = list(ESTRUCTURA_ARCHIVOS.items())
    for idx, (parte, bloques) in enumerate(partes_list):
        print(f"\n📚 Procesando {parte}...")
        
        # Título de parte
        parte_heading = doc.add_heading(parte, level=1)
        parte_heading.paragraph_format.space_before = Pt(6)
        parte_heading.paragraph_format.space_after = Pt(3)
        
        for bloque, archivos in bloques.items():
            print(f"  📖 Procesando {bloque}...")
            
            # Título de bloque
            bloque_heading = doc.add_heading(bloque, level=2)
            bloque_heading.paragraph_format.space_before = Pt(6)
            bloque_heading.paragraph_format.space_after = Pt(3)
            
            # Procesar cada archivo del bloque
            for archivo_relativo in archivos:
                archivo_actual += 1
                archivo_path = BASE_DIR / archivo_relativo.lstrip('./')
                
                if not archivo_path.exists():
                    print(f"    ⚠️  Archivo no encontrado: {archivo_relativo}")
                    p = doc.add_paragraph(f"[ARCHIVO NO ENCONTRADO: {archivo_relativo}]")
                    p.style.font.color.rgb = RGBColor(255, 0, 0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(3)
                    continue
                
                print(f"    📄 [{archivo_actual}/{total_archivos}] {archivo_path.name}")
                
                # Título del capítulo (nombre del archivo sin extensión y formato)
                chapter_title = archivo_path.stem.replace('BLOQUE_', '').replace('_', ' ')
                heading = doc.add_heading(chapter_title, level=3)
                heading.paragraph_format.space_before = Pt(4)
                heading.paragraph_format.space_after = Pt(2)
                
                # Procesar contenido
                content = process_markdown_file(archivo_path)
                if content:
                    process_markdown_content(doc, content, archivo_path.name)
                else:
                    p = doc.add_paragraph("[Error al leer el contenido del archivo]")
                    p.style.font.color.rgb = RGBColor(255, 0, 0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(3)
        
        # Salto de página solo entre partes principales (no después de la última)
        if idx < len(partes_list) - 1:
            doc.add_page_break()
    
    # Guardar documento
    print(f"\n💾 Guardando documento en {OUTPUT_FILE}...")
    doc.save(str(OUTPUT_FILE))
    
    file_size_mb = OUTPUT_FILE.stat().st_size / 1024 / 1024
    print(f"\n✅ Documento generado exitosamente!")
    print(f"   📁 Archivo: {OUTPUT_FILE}")
    print(f"   📊 Tamaño: {file_size_mb:.2f} MB")
    print(f"   📄 Archivos procesados: {archivo_actual}/{total_archivos}")
    print("\n" + "=" * 70)

if __name__ == "__main__":
    main()
