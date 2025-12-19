#!/usr/bin/env python3
"""
Script para convertir archivos Markdown del Manual TES Digital a formato Word (.docx)

Uso:
    python3 convertir_a_word.py [--directorio DIR] [--salida DIR_SALIDA] [--formato docx|xlsx]

Requisitos:
    - pandoc (recomendado) O
    - python-docx y markdown (alternativa)
"""

import os
import sys
import subprocess
import shutil
from pathlib import Path
from typing import Optional, List
import argparse

def check_pandoc() -> bool:
    """Verifica si pandoc está instalado"""
    try:
        result = subprocess.run(['pandoc', '--version'], 
                              capture_output=True, 
                              text=True,
                              timeout=5)
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False

def convert_md_to_docx_pandoc(md_file: Path, docx_file: Path) -> tuple[bool, str]:
    """Convierte un archivo MD a DOCX usando pandoc"""
    try:
        # Comando pandoc para convertir MD a DOCX
        cmd = [
            'pandoc',
            str(md_file),
            '-o', str(docx_file),
            '--from', 'markdown',
            '--to', 'docx',
            '--reference-doc',  # Opcional: usar plantilla personalizada
        ]
        
        # Ejecutar conversión
        result = subprocess.run(cmd, 
                              capture_output=True, 
                              text=True,
                              timeout=30)
        
        if result.returncode == 0:
            return (True, "Convertido correctamente con pandoc")
        else:
            return (False, f"Error pandoc: {result.stderr}")
            
    except subprocess.TimeoutExpired:
        return (False, "Timeout en la conversión")
    except Exception as e:
        return (False, f"Error: {str(e)}")

def convert_md_to_docx_python(md_file: Path, docx_file: Path) -> tuple[bool, str]:
    """Convierte un archivo MD a DOCX usando python-docx (alternativa)"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        # Leer archivo Markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilo por defecto
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        i = 0
        in_list = False
        list_level = 0
        
        while i < len(lines):
            line = lines[i].rstrip('\n\r')
            
            # Línea vacía
            if not line.strip():
                if not in_list:
                    doc.add_paragraph()
                in_list = False
                i += 1
                continue
            
            # Separador horizontal
            if line.strip() == '---':
                doc.add_paragraph('_' * 50)
                i += 1
                continue
            
            # Título nivel 1
            if line.startswith('# '):
                title = line[2:].strip()
                doc.add_heading(title, level=1)
                in_list = False
            # Título nivel 2
            elif line.startswith('## '):
                title = line[3:].strip()
                # Eliminar numeración de secciones (ej: "1.1.1 Objetivo" -> "Objetivo")
                title = re.sub(r'^\d+\.\d+\.\d+\s+', '', title)
                doc.add_heading(title, level=2)
                in_list = False
            # Título nivel 3
            elif line.startswith('### '):
                title = line[4:].strip()
                doc.add_heading(title, level=3)
                in_list = False
            # Lista con viñetas
            elif re.match(r'^[\s]*[-*]\s+', line):
                content = re.sub(r'^[\s]*[-*]\s+', '', line)
                # Manejar indentación (niveles de lista)
                indent_match = re.match(r'^(\s+)', line)
                level = len(indent_match.group(1)) // 2 if indent_match else 0
                
                p = doc.add_paragraph(content, style='List Bullet')
                in_list = True
            # Lista numerada
            elif re.match(r'^\s*\d+\.\s+', line):
                content = re.sub(r'^\s*\d+\.\s+', '', line)
                p = doc.add_paragraph(content, style='List Number')
                in_list = True
            # Texto con formato
            else:
                p = doc.add_paragraph()
                # Procesar negritas, cursivas y código inline
                content = line
                
                # Negritas: **texto**
                parts = re.split(r'(\*\*[^*]+\*\*)', content)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                        # Podría ser cursiva
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Código inline
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        p.add_run(part)
                
                in_list = False
            
            i += 1
        
        # Guardar documento
        doc.save(str(docx_file))
        return (True, "Convertido correctamente con python-docx")
        
    except ImportError as e:
        return (False, f"Biblioteca faltante: {str(e)}. Instala con: pip install python-docx")
    except Exception as e:
        import traceback
        return (False, f"Error en conversión: {str(e)}\n{traceback.format_exc()}")

def convert_directory(source_dir: Path, output_dir: Path, use_pandoc: bool = True) -> dict:
    """Convierte todos los archivos .md de un directorio a .docx"""
    results = {
        'total': 0,
        'exitosos': 0,
        'fallidos': 0,
        'errores': []
    }
    
    # Crear directorio de salida si no existe
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Buscar todos los archivos .md
    md_files = list(source_dir.rglob('*.md'))
    
    # Excluir archivos en _DOCUMENTACION_INTERNA y archivos especiales
    excluded_patterns = ['_DOCUMENTACION_INTERNA', 'MAPA_MAESTRO', 'INFORME', 'ANALISIS', 'ESTANDAR']
    md_files = [f for f in md_files if not any(pattern in str(f) for pattern in excluded_patterns)]
    
    results['total'] = len(md_files)
    
    print(f"\nEncontrados {results['total']} archivos .md para convertir\n")
    
    for md_file in sorted(md_files):
        # Crear ruta relativa para mantener estructura
        relative_path = md_file.relative_to(source_dir)
        
        # Crear estructura de directorios en salida
        docx_file = output_dir / relative_path.with_suffix('.docx')
        docx_file.parent.mkdir(parents=True, exist_ok=True)
        
        print(f"Convirtiendo: {relative_path}...", end=' ')
        
        # Convertir según método disponible
        if use_pandoc:
            success, message = convert_md_to_docx_pandoc(md_file, docx_file)
        else:
            success, message = convert_md_to_docx_python(md_file, docx_file)
        
        if success:
            results['exitosos'] += 1
            print(f"✅")
        else:
            results['fallidos'] += 1
            results['errores'].append({
                'archivo': str(relative_path),
                'error': message
            })
            print(f"❌ {message}")
    
    return results

def main():
    parser = argparse.ArgumentParser(
        description='Convierte archivos Markdown del Manual TES Digital a formato Word (.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  # Convertir todo el directorio TES_Manual_Digital
  python3 convertir_a_word.py
  
  # Especificar directorio fuente y salida
  python3 convertir_a_word.py --directorio ./TES_Manual_Digital --salida ./Manual_Word
  
  # Forzar uso de python-docx (si pandoc no está disponible)
  python3 convertir_a_word.py --no-pandoc
        """
    )
    
    parser.add_argument(
        '--directorio',
        type=str,
        default='.',
        help='Directorio fuente con archivos .md (default: directorio actual)'
    )
    
    parser.add_argument(
        '--salida',
        type=str,
        default='Manual_Word',
        help='Directorio de salida para archivos .docx (default: Manual_Word)'
    )
    
    parser.add_argument(
        '--no-pandoc',
        action='store_true',
        help='Forzar uso de python-docx en lugar de pandoc'
    )
    
    args = parser.parse_args()
    
    # Resolver rutas
    source_dir = Path(args.directorio).resolve()
    output_dir = Path(args.salida).resolve()
    
    if not source_dir.exists():
        print(f"❌ Error: El directorio fuente no existe: {source_dir}")
        sys.exit(1)
    
    # Verificar método de conversión
    use_pandoc = False
    if not args.no_pandoc:
        if check_pandoc():
            use_pandoc = True
            print("✅ Pandoc detectado - usando pandoc para conversión (más fiel al formato original)")
        else:
            print("⚠️  Pandoc no detectado - intentando usar python-docx")
            print("   Para mejor calidad, instala pandoc: sudo apt install pandoc")
    
    # Convertir
    results = convert_directory(source_dir, output_dir, use_pandoc=use_pandoc)
    
    # Resumen
    print(f"\n{'='*60}")
    print(f"RESUMEN DE CONVERSIÓN")
    print(f"{'='*60}")
    print(f"Total de archivos: {results['total']}")
    print(f"✅ Convertidos exitosamente: {results['exitosos']}")
    print(f"❌ Fallidos: {results['fallidos']}")
    
    if results['errores']:
        print(f"\nErrores encontrados:")
        for error in results['errores'][:10]:  # Mostrar máximo 10 errores
            print(f"  - {error['archivo']}: {error['error']}")
        if len(results['errores']) > 10:
            print(f"  ... y {len(results['errores']) - 10} errores más")
    
    print(f"\nArchivos guardados en: {output_dir}")
    print(f"{'='*60}\n")

if __name__ == '__main__':
    main()
