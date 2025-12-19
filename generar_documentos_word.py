#!/usr/bin/env python3
"""
Script para convertir archivos MD de cada módulo en documentos Word (.docx)
Un documento Word por módulo, incluyendo todos sus bloques/temas
"""

import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Mapeo de carpetas a nombres de módulos
MODULOS_MAP = {
    '00_INDICE_Y_NAVEGACION': '00_INDICE_Y_NAVEGACION',
    '01_FUNDAMENTOS_Y_CONCEPTOS': '01_FUNDAMENTOS_Y_CONCEPTOS',
    '02_PROCEDIMIENTOS_BASICOS': '02_PROCEDIMIENTOS_BASICOS',
    '03_SOPORTE_VITAL_BASICO': '03_SOPORTE_VITAL_BASICO',
    '04_MATERIAL_E_INMOVILIZACION': '04_MATERIAL_E_INMOVILIZACION',
    '05_OXIGENOTERAPIA_Y_MATERIAL_SANITARIO': '05_OXIGENOTERAPIA_Y_MATERIAL_SANITARIO',
    '06_PROTOCOLOS_TRANSTELEFONICOS': '06_PROTOCOLOS_TRANSTELEFONICOS',
    '07_FARMACOLOGIA_OPERATIVA': '07_FARMACOLOGIA_OPERATIVA',
    '08_TRANSFERENCIA_Y_TRASLADO': '08_TRANSFERENCIA_Y_TRASLADO',
    '09_TRIAGE_MULTIPLES_VICTIMAS': '09_TRIAGE_MULTIPLES_VICTIMAS',
    '10_CONDUCCION_Y_SEGURIDAD_VIAL': '10_CONDUCCION_Y_SEGURIDAD_VIAL',
    '11_SITUACIONES_ESPECIALES': '11_SITUACIONES_ESPECIALES',
    '12_GESTION_OPERATIVA': '12_GESTION_OPERATIVA',
    'HERRAMIENTAS_Y_CHECKLISTS': 'HERRAMIENTAS_Y_CHECKLISTS',
}

def limpiar_markdown(texto):
    """Limpia y prepara texto Markdown para Word"""
    # Eliminar enlaces pero mantener texto
    texto = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', texto)
    # Eliminar código inline pero mantener texto
    texto = re.sub(r'`([^`]+)`', r'\1', texto)
    return texto

def procesar_markdown_a_docx(doc, contenido_md, archivo_origen):
    """Procesa contenido Markdown y lo agrega al documento Word"""
    lineas = contenido_md.split('\n')
    i = 0
    
    while i < len(lineas):
        linea = lineas[i].strip()
        
        # Saltar líneas vacías
        if not linea:
            i += 1
            continue
        
        # Títulos (# ## ###)
        if linea.startswith('# '):
            doc.add_heading(linea[2:].strip(), level=1)
        elif linea.startswith('## '):
            doc.add_heading(linea[3:].strip(), level=2)
        elif linea.startswith('### '):
            doc.add_heading(linea[4:].strip(), level=3)
        elif linea.startswith('#### '):
            doc.add_heading(linea[5:].strip(), level=4)
        
        # Listas no ordenadas
        elif linea.startswith('- ') or linea.startswith('* '):
            texto = limpiar_markdown(linea[2:].strip())
            p = doc.add_paragraph(texto, style='List Bullet')
            # Procesar líneas siguientes de la lista
            i += 1
            while i < len(lineas) and (lineas[i].strip().startswith('- ') or 
                                       lineas[i].strip().startswith('* ') or
                                       (lineas[i].strip().startswith('  ') and lineas[i].strip())):
                siguiente = lineas[i].strip()
                if siguiente.startswith('- ') or siguiente.startswith('* '):
                    texto = limpiar_markdown(siguiente[2:].strip())
                    doc.add_paragraph(texto, style='List Bullet')
                elif siguiente.startswith('  '):
                    texto = limpiar_markdown(siguiente.strip())
                    if texto:
                        doc.add_paragraph(texto, style='List Bullet 2')
                i += 1
            continue
        
        # Listas ordenadas
        elif re.match(r'^\d+\.\s', linea):
            texto = limpiar_markdown(re.sub(r'^\d+\.\s', '', linea))
            p = doc.add_paragraph(texto, style='List Number')
            i += 1
            while i < len(lineas) and (re.match(r'^\d+\.\s', lineas[i].strip()) or
                                      (lineas[i].strip().startswith('  ') and lineas[i].strip())):
                siguiente = lineas[i].strip()
                if re.match(r'^\d+\.\s', siguiente):
                    texto = limpiar_markdown(re.sub(r'^\d+\.\s', '', siguiente))
                    doc.add_paragraph(texto, style='List Number')
                elif siguiente.startswith('  '):
                    texto = limpiar_markdown(siguiente.strip())
                    if texto:
                        doc.add_paragraph(texto, style='List Number 2')
                i += 1
            continue
        
        # Tablas (básico)
        elif '|' in linea and linea.count('|') >= 2:
            # Detectar inicio de tabla
            filas_tabla = []
            j = i
            while j < len(lineas) and '|' in lineas[j]:
                fila = [celda.strip() for celda in lineas[j].split('|') if celda.strip()]
                if fila and not all(c in '-: ' for c in ''.join(fila)):  # No es separador
                    filas_tabla.append(fila)
                j += 1
            
            if filas_tabla:
                # Crear tabla en Word
                tabla = doc.add_table(rows=len(filas_tabla), cols=len(filas_tabla[0]))
                tabla.style = 'Light Grid Accent 1'
                
                for row_idx, fila in enumerate(filas_tabla):
                    for col_idx, celda in enumerate(fila):
                        if col_idx < len(tabla.rows[row_idx].cells):
                            tabla.rows[row_idx].cells[col_idx].text = limpiar_markdown(celda)
                
                i = j
                continue
        
        # Código bloque (ignorar por ahora o convertir a texto)
        elif linea.startswith('```'):
            # Saltar bloque de código
            i += 1
            while i < len(lineas) and not lineas[i].strip().startswith('```'):
                i += 1
        
        # Separadores
        elif linea.startswith('---') or linea.startswith('***'):
            doc.add_paragraph('─' * 60)
        
        # Párrafo normal
        else:
            texto = limpiar_markdown(linea)
            if texto:
                doc.add_paragraph(texto)
        
        i += 1

def crear_documento_modulo(nombre_modulo, archivos_md, ruta_base, output_dir):
    """Crea un documento Word para un módulo específico"""
    print(f"\n📖 Creando documento: {nombre_modulo}.docx")
    
    # Crear nuevo documento Word
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Portada
    titulo_portada = doc.add_heading('MANUAL TES DIGITAL', 0)
    titulo_portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nombre_modulo_limpio = nombre_modulo.replace('_', ' ').title()
    subtitulo = doc.add_heading(nombre_modulo_limpio, 1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph(f"Documento generado automáticamente desde archivos Markdown")
    fecha = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Fecha: {fecha}")
    
    doc.add_page_break()
    
    # Índice
    doc.add_heading('ÍNDICE DE CONTENIDOS', 1)
    doc.add_paragraph()
    
    for idx, archivo_md in enumerate(archivos_md, 1):
        # Extraer nombre legible
        nombre_tema = archivo_md.replace('BLOQUE_', '').replace('.md', '')
        nombre_tema = re.sub(r'(\d+)_', r'\1. ', nombre_tema)
        nombre_tema = nombre_tema.replace('_', ' ').title()
        doc.add_paragraph(f"{idx}. {nombre_tema}", style='List Number')
    
    doc.add_page_break()
    
    # Procesar cada archivo MD del módulo
    archivos_procesados = 0
    archivos_no_encontrados = []
    
    for archivo_md in archivos_md:
        ruta_archivo = os.path.join(ruta_base, archivo_md)
        
        if not os.path.exists(ruta_archivo):
            print(f"  ⚠️  Archivo no encontrado: {archivo_md}")
            archivos_no_encontrados.append(archivo_md)
            continue
        
        print(f"  ➕ Procesando: {archivo_md}")
        
        # Leer contenido MD
        try:
            with open(ruta_archivo, 'r', encoding='utf-8') as f:
                contenido_md = f.read()
        except Exception as e:
            print(f"  ❌ Error leyendo {archivo_md}: {e}")
            continue
        
        # Extraer título principal (primera línea con #)
        lineas = contenido_md.split('\n')
        titulo_principal = None
        
        for linea in lineas[:10]:  # Buscar en primeras 10 líneas
            if linea.startswith('# '):
                titulo_principal = linea[2:].strip()
                break
        
        if not titulo_principal:
            titulo_principal = archivo_md.replace('.md', '').replace('_', ' ').title()
        
        # Agregar título al documento Word
        doc.add_heading(titulo_principal, level=1)
        
        # Procesar contenido
        procesar_markdown_a_docx(doc, contenido_md, archivo_md)
        
        # Separador entre bloques
        doc.add_paragraph()
        doc.add_paragraph('─' * 60)
        doc.add_paragraph()
        
        archivos_procesados += 1
    
    # Guardar documento
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{nombre_modulo}.docx")
    
    try:
        doc.save(output_path)
        tamaño_kb = os.path.getsize(output_path) / 1024
        print(f"  ✅ Documento guardado: {output_path} ({tamaño_kb:.1f} KB)")
        print(f"  📄 Archivos procesados: {archivos_procesados}/{len(archivos_md)}")
        
        if archivos_no_encontrados:
            print(f"  ⚠️  Archivos no encontrados: {len(archivos_no_encontrados)}")
        
        return output_path, archivos_procesados, archivos_no_encontrados
    except Exception as e:
        print(f"  ❌ Error guardando documento: {e}")
        return None, 0, archivos_no_encontrados

def main():
    """Función principal"""
    print("=" * 70)
    print("GENERADOR DE DOCUMENTOS WORD - MANUAL TES DIGITAL")
    print("=" * 70)
    
    base_path = 'MANUAL_TES_DIGITAL'
    output_dir = 'DOCUMENTOS_WORD'
    
    if not os.path.exists(base_path):
        print(f"❌ Error: La carpeta '{base_path}' no existe.")
        return
    
    # Escanear módulos
    print(f"\n🔍 Explorando módulos en: {base_path}")
    modulos_archivos = {}
    
    for modulo_folder in MODULOS_MAP.keys():
        modulo_path = os.path.join(base_path, modulo_folder)
        if os.path.exists(modulo_path):
            archivos = sorted([f for f in os.listdir(modulo_path) if f.endswith('.md')])
            if archivos:
                modulos_archivos[modulo_folder] = archivos
    
    print(f"\n📚 Módulos encontrados con contenido: {len(modulos_archivos)}")
    for modulo, archivos in sorted(modulos_archivos.items()):
        print(f"  • {modulo}: {len(archivos)} archivos")
    
    if not modulos_archivos:
        print("❌ No se encontraron módulos con archivos .md")
        return
    
    # Generar todos los documentos
    print(f"\n🚀 Generando {len(modulos_archivos)} documentos Word...")
    print("=" * 70)
    
    documentos_generados = []
    total_archivos_procesados = 0
    total_archivos_no_encontrados = []
    
    for modulo_folder in sorted(modulos_archivos.keys()):
        archivos = modulos_archivos[modulo_folder]
        nombre_modulo = MODULOS_MAP[modulo_folder]
        modulo_path = os.path.join(base_path, modulo_folder)
        
        try:
            doc_path, procesados, no_encontrados = crear_documento_modulo(
                nombre_modulo, archivos, modulo_path, output_dir
            )
            
            if doc_path:
                documentos_generados.append(doc_path)
                total_archivos_procesados += procesados
                total_archivos_no_encontrados.extend(no_encontrados)
        except Exception as e:
            print(f"❌ Error procesando {modulo_folder}: {e}")
    
    # Resumen final
    print("\n" + "=" * 70)
    print("📊 RESUMEN DE GENERACIÓN")
    print("=" * 70)
    print(f"✅ Documentos generados: {len(documentos_generados)}")
    print(f"📄 Archivos MD procesados: {total_archivos_procesados}")
    
    if total_archivos_no_encontrados:
        print(f"⚠️  Archivos no encontrados: {len(total_archivos_no_encontrados)}")
    
    print(f"\n📁 Documentos guardados en: {output_dir}/")
    print("\n📋 Documentos generados:")
    for doc in sorted(documentos_generados):
        tamaño = os.path.getsize(doc) / 1024  # KB
        print(f"  • {os.path.basename(doc)} ({tamaño:.1f} KB)")
    
    print("\n🎯 ¡Proceso completado!")

if __name__ == "__main__":
    main()
