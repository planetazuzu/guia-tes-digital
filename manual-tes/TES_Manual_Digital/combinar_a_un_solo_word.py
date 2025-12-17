#!/usr/bin/env python3
"""
Script para combinar todos los archivos .docx del Manual TES Digital en un solo documento

Uso:
    python3 combinar_a_un_solo_word.py [--directorio DIR] [--salida archivo.docx]
"""

import os
import sys
from pathlib import Path
from typing import List
import argparse

def combine_docx_files(docx_files: List[Path], output_file: Path) -> tuple[bool, str]:
    """Combina múltiples archivos .docx en uno solo"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.section import WD_SECTION
        
        # Crear nuevo documento
        master_doc = Document()
        
        # Configurar estilo por defecto
        style = master_doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        print(f"\nCombinando {len(docx_files)} archivos en un solo documento...\n")
        
        for i, docx_file in enumerate(docx_files, 1):
            print(f"[{i}/{len(docx_files)}] Procesando: {docx_file.name}...", end=' ')
            
            try:
                # Abrir documento fuente
                source_doc = Document(str(docx_file))
                
                # Agregar título del capítulo (nombre del archivo sin extensión)
                chapter_title = docx_file.stem.replace('_', ' ').replace('BLOQUE', 'BLOQUE')
                
                # Añadir salto de página antes de cada capítulo (excepto el primero)
                if i > 1:
                    master_doc.add_page_break()
                
                # Añadir título del capítulo
                heading = master_doc.add_heading(chapter_title, level=1)
                
                # Copiar contenido del documento fuente
                for element in source_doc.element.body:
                    # Saltar el primer elemento si es el título (ya lo agregamos)
                    if element.tag.endswith('p'):
                        para = source_doc.paragraphs[0] if len(source_doc.paragraphs) > 0 else None
                        if para and para.style.name.startswith('Heading 1'):
                            continue  # Saltar el título original
                    
                    # Copiar el elemento
                    master_doc.element.body.append(element)
                
                print("✅")
                
            except Exception as e:
                print(f"⚠️  Error: {str(e)}")
                # Continuar con el siguiente archivo
                continue
        
        # Guardar documento combinado
        master_doc.save(str(output_file))
        return (True, f"Documento combinado guardado en: {output_file}")
        
    except ImportError:
        return (False, "Biblioteca faltante: python-docx. Instala con: pip install python-docx")
    except Exception as e:
        return (False, f"Error: {str(e)}")

def combine_docx_files_advanced(docx_files: List[Path], output_file: Path) -> tuple[bool, str]:
    """Combina archivos .docx de forma más avanzada, preservando mejor el formato"""
    try:
        from docx import Document
        from docx.shared import Pt
        import xml.etree.ElementTree as ET
        
        print(f"\nCombinando {len(docx_files)} archivos en un solo documento...\n")
        
        # Crear documento maestro
        master_doc = Document()
        
        # Configurar estilo
        style = master_doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        for i, docx_file in enumerate(docx_files, 1):
            print(f"[{i}/{len(docx_files)}] Procesando: {docx_file.name}...", end=' ')
            
            try:
                source_doc = Document(str(docx_file))
                
                # Añadir salto de página antes de cada capítulo (excepto el primero)
                if i > 1:
                    master_doc.add_page_break()
                    # Añadir línea separadora
                    master_doc.add_paragraph('_' * 80)
                    master_doc.add_paragraph()
                
                # Copiar todos los párrafos del documento fuente
                first_heading_skipped = False
                
                for para in source_doc.paragraphs:
                    # Saltar el primer heading (título principal) del documento fuente
                    if not first_heading_skipped and para.style.name.startswith('Heading'):
                        first_heading_skipped = True
                        # Añadir el título del capítulo como heading de nivel 1
                        chapter_title = docx_file.stem.replace('_', ' ')
                        master_doc.add_heading(chapter_title, level=1)
                        continue
                    
                    # Copiar el párrafo
                    new_para = master_doc.add_paragraph()
                    
                    # Copiar el texto y formato
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                    
                    # Aplicar el estilo del párrafo original
                    if para.style.name != 'Normal':
                        try:
                            new_para.style = para.style.name
                        except:
                            pass  # Si el estilo no existe, mantener Normal
                
                # Copiar tablas si las hay
                for table in source_doc.tables:
                    # Crear nueva tabla con las mismas dimensiones
                    new_table = master_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style
                    
                    # Copiar contenido de celdas
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                
                print("✅")
                
            except Exception as e:
                print(f"⚠️  Error: {str(e)}")
                continue
        
        # Guardar documento final
        master_doc.save(str(output_file))
        return (True, f"Documento combinado guardado exitosamente en: {output_file}")
        
    except ImportError:
        return (False, "Biblioteca faltante: python-docx. Instala con: pip install python-docx")
    except Exception as e:
        import traceback
        return (False, f"Error: {str(e)}\n{traceback.format_exc()}")

def get_sorted_docx_files(source_dir: Path) -> List[Path]:
    """Obtiene lista de archivos .docx ordenados por bloque y capítulo"""
    excluded_patterns = ['INSTALACION', 'README', 'LEEME']
    
    # Buscar todos los archivos .docx
    docx_files = list(source_dir.rglob('*.docx'))
    
    # Filtrar archivos excluidos
    docx_files = [f for f in docx_files if not any(pattern in f.name for pattern in excluded_patterns)]
    
    # Ordenar por ruta (esto mantiene el orden de bloques y capítulos)
    docx_files.sort(key=lambda x: str(x))
    
    return docx_files

def main():
    parser = argparse.ArgumentParser(
        description='Combina todos los archivos .docx del Manual TES Digital en un solo documento',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  # Combinar archivos del directorio Manual_Word
  python3 combinar_a_un_solo_word.py --directorio ./Manual_Word
  
  # Especificar directorio y nombre de salida
  python3 combinar_a_un_solo_word.py --directorio ./Manual_Word --salida Manual_TES_Completo.docx
        """
    )
    
    parser.add_argument(
        '--directorio',
        type=str,
        default='./Manual_Word',
        help='Directorio con archivos .docx a combinar (default: ./Manual_Word)'
    )
    
    parser.add_argument(
        '--salida',
        type=str,
        default='Manual_TES_Digital_Completo.docx',
        help='Nombre del archivo de salida (default: Manual_TES_Digital_Completo.docx)'
    )
    
    args = parser.parse_args()
    
    # Resolver rutas
    source_dir = Path(args.directorio).resolve()
    output_file = Path(args.salida).resolve()
    
    if not source_dir.exists():
        print(f"❌ Error: El directorio fuente no existe: {source_dir}")
        sys.exit(1)
    
    # Obtener lista de archivos ordenados
    docx_files = get_sorted_docx_files(source_dir)
    
    if not docx_files:
        print(f"❌ Error: No se encontraron archivos .docx en {source_dir}")
        sys.exit(1)
    
    print(f"Encontrados {len(docx_files)} archivos .docx para combinar")
    
    # Combinar archivos
    success, message = combine_docx_files_advanced(docx_files, output_file)
    
    if success:
        print(f"\n{'='*60}")
        print("✅ COMBINACIÓN COMPLETADA")
        print(f"{'='*60}")
        print(message)
        print(f"Archivos combinados: {len(docx_files)}")
        print(f"Tamaño del archivo: {output_file.stat().st_size / 1024 / 1024:.2f} MB")
        print(f"{'='*60}\n")
    else:
        print(f"\n❌ Error: {message}\n")
        sys.exit(1)

if __name__ == '__main__':
    main()
